import os
import re
from docx import Document
from models import SessionLocal, Customer, FeeType, init_db
from datetime import date

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "invoice_templates")

def extract_money(text):
    # Find all money patterns like $150 or $1,200.50
    matches = re.findall(r'\$\s?([0-9,]+(?:\.[0-9]{2})?)', text)
    if matches:
        # Return the last one found in the line, assuming it's the total for that line
        return float(matches[-1].replace(',', ''))
    return 0.0

def seed_customers():
    print("Initializing DB...")
    init_db()
    session = SessionLocal()
    
    # Ensure Management Fee exists
    if not session.query(FeeType).filter_by(name="Management Fee").first():
        session.add(FeeType(name="Management Fee"))
        session.commit()

    files = [f for f in os.listdir(TEMPLATE_DIR) if f.lower().endswith('.docx')]
    
    count = 0
    for f in files:
        if f == "base_invoice_template.docx" or f.startswith("~"):
            continue
            
        path = os.path.join(TEMPLATE_DIR, f)
        print(f"Processing {f}...")
        
        try:
            doc = Document(path)
            
            name = ""
            address = ""
            rate = 0.0
            cadence = "monthly" # Default
            fee_type = "Management Fee"
            
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                
                # Extract Name
                if text.upper().startswith("TO:"):
                    # Try to get name from same line
                    parts = text.split(":", 1)
                    if len(parts) > 1 and parts[1].strip():
                        name = parts[1].strip()
                    else:
                        # Name might be on next line, but simple parsing is hard here.
                        # Let's assume it's on the same line for now based on analysis.
                        pass
                
                # Extract Address
                if text.upper().startswith("FOR:"):
                    parts = text.split(":", 1)
                    if len(parts) > 1 and parts[1].strip():
                        address = parts[1].strip()
                
                # Extract Email
                email_matches = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
                for email_match in email_matches:
                    # Filter out sender emails
                    if "linda" not in email_match.lower() and "stonegate" not in email_match.lower():
                         # Assume the first non-sender email is the customer's
                         if not email or email == "change@me.com":
                             email = email_match

                # Extract Rate & Cadence from management line
                if "management" in text.lower() or "quarter" in text.lower():
                    line_rate = extract_money(text)
                    if line_rate > 0:
                        rate = line_rate
                    
                    if "quarter" in text.lower():
                        cadence = "quarterly"
                    elif "year" in text.lower() or "annual" in text.lower():
                        cadence = "yearly"
            
            # Fallback for rate if not found in management line
            if rate == 0:
                for p in doc.paragraphs:
                    if "Total due" in p.text:
                        rate = extract_money(p.text)
                        break

            if name and address:
                # Check if exists
                existing = session.query(Customer).filter(Customer.name == name).first()
                if not existing:
                    print(f"  -> Adding {name} ({address}) - ${rate} {cadence}")
                    c = Customer(
                        name=name,
                        email=email if email else "change@me.com",
                        property_address=address,
                        rate=rate,
                        cadence=cadence,
                        fee_type=fee_type,
                        next_bill_date=date.today()
                    )
                    session.add(c)
                    count += 1
                else:
                    # Update email if missing
                    if email and (not existing.email or existing.email == "change@me.com"):
                        print(f"  -> Updating email for {name}: {email}")
                        existing.email = email
                        session.add(existing)
                        count += 1
                    else:
                        print(f"  -> Skipping {name} (already exists)")
            else:
                print(f"  -> Could not extract Name or Address from {f}")

        except Exception as e:
            print(f"  -> Error processing {f}: {e}")

    session.commit()
    session.close()
    print(f"Done! Added {count} new customers.")

if __name__ == "__main__":
    from datetime import date
    seed_customers()
